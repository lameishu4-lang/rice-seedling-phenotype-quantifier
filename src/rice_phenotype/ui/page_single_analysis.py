# -*- coding: utf-8 -*-

from datetime import datetime
from pathlib import Path

import cv2
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PySide6.QtCore import Qt
from PySide6.QtGui import QPixmap
from PySide6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QLabel,
    QPushButton,
    QFileDialog,
    QFrame,
    QMessageBox,
    QDoubleSpinBox,
)

from rice_phenotype import __version__
from rice_phenotype.core.segmentation import (
    SeedlingSegmenter,
    SegmentationConfig,
    SegmentationResult,
)
from rice_phenotype.core.calibration import ScaleCalibrator
from rice_phenotype.core.metrics import PhenotypeCalculator, PhenotypeMetrics
from rice_phenotype.storage.database import RecordRepository
from rice_phenotype.utils.image_qt import ndarray_to_pixmap
from rice_phenotype.utils.paths import image_output_dir, output_dir


class SingleAnalysisPage(QWidget):
    def __init__(self):
        super().__init__()

        self.current_image_path: Path | None = None
        self.current_image_rgb: np.ndarray | None = None
        self.current_segmentation: SegmentationResult | None = None
        self.current_metrics: PhenotypeMetrics | None = None
        self.current_config: SegmentationConfig | None = None

        self.record_saved = False
        self.report_exported = False
        self.saved_record_id: int | None = None

        self.last_mask_path: Path | None = None
        self.last_overlay_path: Path | None = None

        self.segmenter = SeedlingSegmenter()
        self.calibrator = ScaleCalibrator()
        self.calculator = PhenotypeCalculator()
        self.repository = RecordRepository()

        self._build_ui()
        self._reset_to_initial_state()

    def _build_ui(self) -> None:
        root_layout = QVBoxLayout(self)
        root_layout.setContentsMargins(36, 32, 36, 32)
        root_layout.setSpacing(18)

        title = QLabel("单张图像分析")
        title.setObjectName("PageTitle")
        root_layout.addWidget(title)

        desc = QLabel(
            "导入水稻秧苗图像，完成绿色植株区域分割、比例尺换算、二维表型指标计算、记录保存和报告导出。"
        )
        desc.setObjectName("PageDesc")
        desc.setWordWrap(True)
        root_layout.addWidget(desc)

        toolbar = QHBoxLayout()
        toolbar.setSpacing(12)

        self.btn_import = QPushButton("导入图像")
        self.btn_import.setObjectName("PrimaryButton")
        self.btn_import.clicked.connect(self.import_image)

        self.btn_segment = QPushButton("执行分割")
        self.btn_segment.setObjectName("SecondaryButton")
        self.btn_segment.clicked.connect(self.run_segmentation)

        self.btn_calculate = QPushButton("计算指标")
        self.btn_calculate.setObjectName("SecondaryButton")
        self.btn_calculate.clicked.connect(self.calculate_metrics)

        self.btn_save = QPushButton("保存记录")
        self.btn_save.setObjectName("SecondaryButton")
        self.btn_save.clicked.connect(self.save_record)

        self.btn_report = QPushButton("导出报告")
        self.btn_report.setObjectName("SecondaryButton")
        self.btn_report.clicked.connect(self.export_report)

        toolbar.addWidget(self.btn_import)
        toolbar.addWidget(self.btn_segment)
        toolbar.addWidget(self.btn_calculate)
        toolbar.addWidget(self.btn_save)
        toolbar.addWidget(self.btn_report)
        toolbar.addStretch()

        root_layout.addLayout(toolbar)

        scale_card = QFrame()
        scale_card.setObjectName("Card")
        scale_layout = QHBoxLayout(scale_card)
        scale_layout.setContentsMargins(18, 12, 18, 12)
        scale_layout.setSpacing(12)

        scale_title = QLabel("比例尺设置：")
        scale_title.setStyleSheet("font-size: 14px; font-weight: 700; color: #111827;")
        scale_layout.addWidget(scale_title)

        self.scale_spin = QDoubleSpinBox()
        self.scale_spin.setDecimals(5)
        self.scale_spin.setMinimum(0.00001)
        self.scale_spin.setMaximum(10.0)
        self.scale_spin.setSingleStep(0.001)
        self.scale_spin.setValue(0.05000)
        self.scale_spin.setSuffix(" cm/pixel")
        self.scale_spin.setFixedWidth(180)
        scale_layout.addWidget(self.scale_spin)

        scale_note = QLabel("说明：当前采用手动比例尺，结果用于二维图像辅助量化。")
        scale_note.setStyleSheet("font-size: 13px; color: #4B5563;")
        scale_layout.addWidget(scale_note)
        scale_layout.addStretch()

        root_layout.addWidget(scale_card)

        content_layout = QHBoxLayout()
        content_layout.setSpacing(16)

        self.original_panel = self._create_image_panel("原始图像")
        self.mask_panel = self._create_image_panel("分割掩膜")
        self.overlay_panel = self._create_image_panel("叠加结果")

        content_layout.addWidget(self.original_panel["card"])
        content_layout.addWidget(self.mask_panel["card"])
        content_layout.addWidget(self.overlay_panel["card"])

        root_layout.addLayout(content_layout, stretch=1)

        metrics_card = QFrame()
        metrics_card.setObjectName("Card")
        metrics_layout = QVBoxLayout(metrics_card)
        metrics_layout.setContentsMargins(20, 16, 20, 16)

        metric_title = QLabel("分析信息与表型指标")
        metric_title.setStyleSheet("font-size: 17px; font-weight: 700; color: #111827;")
        metrics_layout.addWidget(metric_title)

        self.metrics_label = QLabel()
        self.metrics_label.setWordWrap(True)
        self.metrics_label.setStyleSheet("font-size: 14px; color: #374151; line-height: 1.6;")
        metrics_layout.addWidget(self.metrics_label)

        root_layout.addWidget(metrics_card)

    def _create_image_panel(self, title: str) -> dict:
        card = QFrame()
        card.setObjectName("Card")

        layout = QVBoxLayout(card)
        layout.setContentsMargins(16, 14, 16, 16)
        layout.setSpacing(10)

        title_label = QLabel(title)
        title_label.setStyleSheet("font-size: 15px; font-weight: 700; color: #111827;")
        layout.addWidget(title_label)

        image_label = QLabel("暂无图像")
        image_label.setAlignment(Qt.AlignCenter)
        image_label.setMinimumHeight(320)
        image_label.setStyleSheet(
            """
            QLabel {
                background-color: #F9FAFB;
                border: 1px dashed #CBD5E1;
                border-radius: 10px;
                color: #6B7280;
            }
            """
        )
        layout.addWidget(image_label, stretch=1)

        return {
            "card": card,
            "label": image_label,
        }

    def _reset_to_initial_state(self) -> None:
        self.current_image_path = None
        self.current_image_rgb = None
        self.current_segmentation = None
        self.current_metrics = None
        self.current_config = None

        self.record_saved = False
        self.report_exported = False
        self.saved_record_id = None
        self.last_mask_path = None
        self.last_overlay_path = None

        self.btn_segment.setEnabled(False)
        self.btn_calculate.setEnabled(False)
        self.btn_save.setEnabled(False)
        self.btn_report.setEnabled(False)

        self.metrics_label.setText(
            "尚未导入图像。\n\n"
            "操作流程：导入图像 → 执行分割 → 计算指标 → 保存记录 / 导出报告。\n\n"
            "说明：保存记录和导出报告必须在表型指标计算完成后才能执行，且同一轮分析结果只能各执行一次。"
        )

    def _reset_after_import(self) -> None:
        self.current_segmentation = None
        self.current_metrics = None
        self.current_config = None

        self.record_saved = False
        self.report_exported = False
        self.saved_record_id = None
        self.last_mask_path = None
        self.last_overlay_path = None

        self.btn_segment.setEnabled(True)
        self.btn_calculate.setEnabled(False)
        self.btn_save.setEnabled(False)
        self.btn_report.setEnabled(False)

    def _reset_after_segmentation(self) -> None:
        self.current_metrics = None

        self.record_saved = False
        self.report_exported = False
        self.saved_record_id = None
        self.last_mask_path = None
        self.last_overlay_path = None

        self.btn_calculate.setEnabled(True)
        self.btn_save.setEnabled(False)
        self.btn_report.setEnabled(False)

    def _reset_after_metric_calculation(self) -> None:
        self.record_saved = False
        self.report_exported = False
        self.saved_record_id = None
        self.last_mask_path = None
        self.last_overlay_path = None

        self.btn_save.setEnabled(True)
        self.btn_report.setEnabled(True)

    def import_image(self) -> None:
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择水稻秧苗图像",
            "",
            "Image Files (*.jpg *.jpeg *.png *.bmp)",
        )

        if not file_path:
            return

        path = Path(file_path)

        image_bgr = cv2.imdecode(
            np.fromfile(str(path), dtype=np.uint8),
            cv2.IMREAD_COLOR,
        )

        if image_bgr is None:
            QMessageBox.warning(self, "图像读取失败", "无法读取该图像文件，请检查文件格式或路径。")
            return

        image_rgb = cv2.cvtColor(image_bgr, cv2.COLOR_BGR2RGB)

        self.current_image_path = path
        self.current_image_rgb = image_rgb

        self._show_ndarray_image(self.original_panel["label"], image_rgb)

        self.mask_panel["label"].setPixmap(QPixmap())
        self.mask_panel["label"].setText("待执行分割")

        self.overlay_panel["label"].setPixmap(QPixmap())
        self.overlay_panel["label"].setText("待执行分割")

        self._reset_after_import()

        height, width = image_rgb.shape[:2]

        self.metrics_label.setText(
            f"已导入图像：{path.name}\n"
            f"图像尺寸：{width} × {height} px\n"
            f"文件路径：{path}\n\n"
            "下一步：点击“执行分割”生成秧苗区域掩膜。"
        )

    def run_segmentation(self) -> None:
        if self.current_image_rgb is None:
            QMessageBox.information(self, "提示", "请先导入图像。")
            return

        config = SegmentationConfig(
            hsv_lower=(35, 40, 40),
            hsv_upper=(85, 255, 255),
            exg_threshold=30,
            min_area=300,
            kernel_size=5,
            method="HSV",
        )

        try:
            result = self.segmenter.segment(self.current_image_rgb, config)
        except Exception as exc:
            QMessageBox.critical(self, "分割失败", f"执行分割时发生错误：\n{exc}")
            return

        self.current_segmentation = result
        self.current_config = config

        self._show_ndarray_image(self.mask_panel["label"], result.mask)
        self._show_ndarray_image(self.overlay_panel["label"], result.overlay)

        x, y, w, h = result.bbox

        if result.success:
            self._reset_after_segmentation()

            self.metrics_label.setText(
                "分割完成。\n\n"
                f"分割方法：{config.method}\n"
                f"HSV 下限：{config.hsv_lower}\n"
                f"HSV 上限：{config.hsv_upper}\n"
                f"最小连通域面积：{config.min_area} px\n\n"
                f"图像有效面积：{result.valid_area_px} px\n"
                f"秧苗掩膜面积：{result.plant_area_px} px\n"
                f"外接矩形：x={x}, y={y}, w={w}, h={h}\n\n"
                "下一步：确认比例尺后，点击“计算指标”。"
            )
        else:
            self.btn_calculate.setEnabled(False)
            self.btn_save.setEnabled(False)
            self.btn_report.setEnabled(False)

            self.metrics_label.setText(
                "分割未检测到有效区域。\n\n"
                f"提示信息：{result.message}\n\n"
                "建议：更换图像、调整拍摄背景，或后续在参数设置页中调整 HSV / ExG 阈值。"
            )

    def calculate_metrics(self) -> None:
        if self.current_image_rgb is None:
            QMessageBox.information(self, "提示", "请先导入图像。")
            return

        if self.current_segmentation is None or not self.current_segmentation.success:
            QMessageBox.information(self, "提示", "请先完成有效分割。")
            return

        cm_per_pixel = float(self.scale_spin.value())

        try:
            self.calibrator.validate_cm_per_pixel(cm_per_pixel)

            metrics = self.calculator.calculate(
                image_rgb=self.current_image_rgb,
                mask=self.current_segmentation.mask,
                cm_per_pixel=cm_per_pixel,
                valid_area_px=self.current_segmentation.valid_area_px,
            )
        except Exception as exc:
            QMessageBox.critical(self, "指标计算失败", f"计算表型指标时发生错误：\n{exc}")
            return

        self.current_metrics = metrics
        self._reset_after_metric_calculation()
        self.metrics_label.setText(self._format_metrics(metrics, cm_per_pixel))

    def save_record(self) -> None:
        if self.record_saved:
            QMessageBox.information(self, "提示", "当前分析结果已经保存过记录。")
            self.btn_save.setEnabled(False)
            return

        if not self._has_complete_analysis():
            return

        try:
            mask_path, overlay_path = self._ensure_analysis_images_saved()
            save_time = datetime.now()

            height, width = self.current_image_rgb.shape[:2]
            metrics = self.current_metrics
            config = self.current_config
            cm_per_pixel = float(self.scale_spin.value())

            record = {
                "sample_name": self.current_image_path.name,
                "image_path": str(self.current_image_path),
                "mask_path": str(mask_path),
                "overlay_path": str(overlay_path),
                "analysis_time": save_time.strftime("%Y-%m-%d %H:%M:%S"),
                "image_width": width,
                "image_height": height,
                "cm_per_pixel": cm_per_pixel,
                "segmentation_method": config.method,
                "hsv_lower": str(config.hsv_lower),
                "hsv_upper": str(config.hsv_upper),
                "exg_threshold": config.exg_threshold,
                "plant_height_px": metrics.plant_height_px,
                "plant_height_cm": metrics.plant_height_cm,
                "canopy_width_px": metrics.canopy_width_px,
                "canopy_width_cm": metrics.canopy_width_cm,
                "projected_area_px": metrics.projected_area_px,
                "projected_area_cm2": metrics.projected_area_cm2,
                "green_coverage": metrics.green_coverage,
                "exg_mean": metrics.exg_mean,
                "green_ratio": metrics.green_ratio,
                "bbox_fill_ratio": metrics.bbox_fill_ratio,
                "growth_score": metrics.growth_score,
                "note": "",
                "software_version": __version__,
            }

            record_id = self.repository.insert_record(record)
        except Exception as exc:
            QMessageBox.critical(self, "保存失败", f"保存分析记录时发生错误：\n{exc}")
            return

        self.saved_record_id = record_id
        self.record_saved = True
        self.btn_save.setEnabled(False)

        QMessageBox.information(
            self,
            "保存成功",
            f"分析记录已保存。\n数据库记录ID：{record_id}\n\n"
            "说明：数据库记录ID用于内部追踪，不等同于历史记录表中的连续显示序号。"
        )

    def export_report(self) -> None:
        if self.report_exported:
            QMessageBox.information(self, "提示", "当前分析结果已经导出过报告。")
            self.btn_report.setEnabled(False)
            return

        if not self._has_complete_analysis():
            return

        report_folder = output_dir() / "reports"
        report_folder.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = self.current_image_path.stem if self.current_image_path else "sample"
        default_path = report_folder / f"{base_name}_{timestamp}_report.docx"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出单样本 Word 报告",
            str(default_path),
            "Word Documents (*.docx)",
        )

        if not file_path:
            return

        output_path = Path(file_path)

        if output_path.suffix.lower() != ".docx":
            output_path = output_path.with_suffix(".docx")

        try:
            mask_path, overlay_path = self._ensure_analysis_images_saved()
            self._create_word_report(
                output_path=output_path,
                mask_path=mask_path,
                overlay_path=overlay_path,
            )
        except Exception as exc:
            QMessageBox.critical(self, "导出失败", f"导出 Word 报告时发生错误：\n{exc}")
            return

        self.report_exported = True
        self.btn_report.setEnabled(False)

        QMessageBox.information(
            self,
            "导出成功",
            f"单样本图文报告已导出：\n{output_path}"
        )

    def _has_complete_analysis(self) -> bool:
        if self.current_image_path is None:
            QMessageBox.information(self, "提示", "请先导入图像。")
            return False

        if self.current_image_rgb is None:
            QMessageBox.information(self, "提示", "当前图像为空。")
            return False

        if self.current_segmentation is None or not self.current_segmentation.success:
            QMessageBox.information(self, "提示", "请先完成有效分割。")
            return False

        if self.current_metrics is None:
            QMessageBox.information(self, "提示", "请先计算表型指标。")
            return False

        if self.current_config is None:
            QMessageBox.information(self, "提示", "缺少分割参数。")
            return False

        return True

    def _ensure_analysis_images_saved(self) -> tuple[Path, Path]:
        if self.last_mask_path is not None and self.last_overlay_path is not None:
            if self.last_mask_path.exists() and self.last_overlay_path.exists():
                return self.last_mask_path, self.last_overlay_path

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = self.current_image_path.stem

        output_folder = image_output_dir()

        mask_path = output_folder / f"{base_name}_{timestamp}_mask.png"
        overlay_path = output_folder / f"{base_name}_{timestamp}_overlay.png"

        self._save_mask_image(mask_path, self.current_segmentation.mask)
        self._save_rgb_image(overlay_path, self.current_segmentation.overlay)

        self.last_mask_path = mask_path
        self.last_overlay_path = overlay_path

        return mask_path, overlay_path

    def _create_word_report(
        self,
        output_path: Path,
        mask_path: Path,
        overlay_path: Path,
    ) -> None:
        metrics = self.current_metrics
        config = self.current_config
        cm_per_pixel = float(self.scale_spin.value())

        doc = Document()

        title = doc.add_heading("水稻秧苗图像表型量化分析报告", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"软件版本：V{__version__}")
        doc.add_paragraph(f"样本名称：{self.current_image_path.name}")
        doc.add_paragraph(f"原始图像路径：{self.current_image_path}")

        doc.add_heading("一、分析参数", level=1)
        doc.add_paragraph(f"比例尺：{cm_per_pixel:.5f} cm/pixel")
        doc.add_paragraph(f"分割方法：{config.method}")
        doc.add_paragraph(f"HSV 下限：{config.hsv_lower}")
        doc.add_paragraph(f"HSV 上限：{config.hsv_upper}")
        doc.add_paragraph(f"ExG 阈值：{config.exg_threshold}")
        doc.add_paragraph(f"最小连通域面积：{config.min_area} px")

        doc.add_heading("二、图像结果", level=1)

        if self.current_image_path.exists():
            doc.add_paragraph("原始图像：")
            doc.add_picture(str(self.current_image_path), width=Inches(5.5))

        if mask_path.exists():
            doc.add_paragraph("分割掩膜图：")
            doc.add_picture(str(mask_path), width=Inches(5.5))

        if overlay_path.exists():
            doc.add_paragraph("叠加结果图：")
            doc.add_picture(str(overlay_path), width=Inches(5.5))

        doc.add_heading("三、表型指标", level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "指标类别"
        header_cells[1].text = "指标名称"
        header_cells[2].text = "计算结果"

        rows = [
            ("形态指标", "株高估算", f"{metrics.plant_height_cm:.2f} cm ({metrics.plant_height_px:.0f} px)"),
            ("形态指标", "冠幅估算", f"{metrics.canopy_width_cm:.2f} cm ({metrics.canopy_width_px:.0f} px)"),
            ("形态指标", "投影面积", f"{metrics.projected_area_cm2:.2f} cm² ({metrics.projected_area_px:.0f} px)"),
            ("形态指标", "绿色覆盖率", f"{metrics.green_coverage * 100:.2f}%"),
            ("颜色指标", "ExG 叶色指数均值", f"{metrics.exg_mean:.2f}"),
            ("颜色指标", "Green Ratio", f"{metrics.green_ratio:.4f}"),
            ("内部评分", "外接矩形填充率", f"{metrics.bbox_fill_ratio * 100:.2f}%"),
            ("内部评分", "长势评分", f"{metrics.growth_score:.2f} / 100"),
        ]

        for category, name, value in rows:
            cells = table.add_row().cells
            cells[0].text = category
            cells[1].text = name
            cells[2].text = value

        doc.add_heading("四、结果说明", level=1)
        doc.add_paragraph(
            "本报告结果基于二维图像分割、像素统计和比例尺换算生成，仅用于样本记录、"
            "教学演示和科研辅助整理，不作为田间生产决策、农学诊断或病害判断的唯一依据。"
        )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    def _format_metrics(
        self,
        metrics: PhenotypeMetrics,
        cm_per_pixel: float,
    ) -> str:
        return (
            "表型指标计算完成。\n\n"
            f"当前比例尺：{cm_per_pixel:.5f} cm/pixel\n\n"
            "【形态指标】\n"
            f"株高估算：{metrics.plant_height_cm:.2f} cm "
            f"({metrics.plant_height_px:.0f} px)\n"
            f"冠幅估算：{metrics.canopy_width_cm:.2f} cm "
            f"({metrics.canopy_width_px:.0f} px)\n"
            f"投影面积：{metrics.projected_area_cm2:.2f} cm² "
            f"({metrics.projected_area_px:.0f} px)\n"
            f"绿色覆盖率：{metrics.green_coverage * 100:.2f}%\n\n"
            "【颜色指标】\n"
            f"ExG 叶色指数均值：{metrics.exg_mean:.2f}\n"
            f"Green Ratio：{metrics.green_ratio:.4f}\n\n"
            "【软件内部评分】\n"
            f"掩膜外接矩形填充率：{metrics.bbox_fill_ratio * 100:.2f}%\n"
            f"长势评分：{metrics.growth_score:.2f} / 100\n\n"
            "说明：以上结果基于二维图像分割和比例尺换算，仅用于样本记录、"
            "教学演示和科研辅助整理，不作为农学诊断或生产决策依据。\n\n"
            "下一步：可点击“保存记录”写入本地数据库，或点击“导出报告”生成单样本 Word 报告。"
        )

    def _show_ndarray_image(self, label: QLabel, image: np.ndarray) -> None:
        pixmap = ndarray_to_pixmap(image)

        target_width = max(label.width(), 320)
        target_height = max(label.height(), 260)

        scaled = pixmap.scaled(
            target_width,
            target_height,
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation,
        )

        label.setPixmap(scaled)
        label.setText("")

    @staticmethod
    def _save_mask_image(path: Path, mask: np.ndarray) -> None:
        success, encoded = cv2.imencode(".png", mask)

        if not success:
            raise RuntimeError("掩膜图编码失败。")

        encoded.tofile(str(path))

    @staticmethod
    def _save_rgb_image(path: Path, image_rgb: np.ndarray) -> None:
        image_bgr = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2BGR)

        success, encoded = cv2.imencode(".png", image_bgr)

        if not success:
            raise RuntimeError("图像编码失败。")

        encoded.tofile(str(path))